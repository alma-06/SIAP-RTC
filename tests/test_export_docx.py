from docx import Document

from app.dashboard.fact_sheet import build_fact_sheet
from app.dashboard.model import DashboardPeriod
from app.export.contracts import build_export_payload
from app.export.docx import export_docx


def test_export_docx_creates_executive_document(tmp_path) -> None:
    period = DashboardPeriod("2026-Q2", 100, 10, 5, 80, 5, 0.80, 0.05, 0.10, 0.05, True, (), "EV-01")
    payload = build_export_payload(build_fact_sheet(period))
    output = export_docx(payload, tmp_path / "siap_rtc.docx")
    document = Document(output)
    text = "\n".join(p.text for p in document.paragraphs)
    assert "SIAP-RTC" in text
    assert "2026-Q2" in text
    assert "Identificador de evidencia: EV-01" in text
    assert len(document.tables) == 2
