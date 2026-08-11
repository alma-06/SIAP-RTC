from __future__ import annotations

from dataclasses import asdict
from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.export.executive_package import build_executive_summary
from app.pipeline.integrated import IntegratedPipelineResult


def write_docx_report(result: IntegratedPipelineResult, path: str | Path) -> None:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    summary = build_executive_summary(result)
    document.add_heading("Informe Ejecutivo — SIAP-RTC", level=0)
    document.add_paragraph(
        "Documento generado a partir del resultado integrado del procesamiento RTC, "
        "conservando trazabilidad hacia las fuentes y advertencias metodológicas."
    )

    document.add_heading("1. Objetivo", level=1)
    document.add_paragraph(
        "Procesar, depurar, consolidar y documentar información de pautas RTC correspondiente "
        "a la Cámara de Senadores, manteniendo trazabilidad de los registros y de los resultados derivados."
    )

    document.add_heading("2. Resultados del procesamiento", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Indicador"
    table.rows[0].cells[1].text = "Resultado"
    for key in ("records_ingested", "records_normalized", "records_kept", "duplicates_removed"):
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = str(summary[key])

    if result.indicators is not None:
        document.add_heading("3. Indicadores de conciliación", level=1)
        for key, value in asdict(result.indicators).items():
            document.add_paragraph(f"{key}: {value}", style="List Bullet")

    document.add_heading("4. Fuentes y evidencia", level=1)
    for source in result.evidence.sources:
        document.add_paragraph(
            f"{source.filename} — {source.period} — SHA-256: {source.sha256}",
            style="List Bullet",
        )

    document.add_heading("5. Consideraciones metodológicas", level=1)
    for warning in result.evidence.warnings:
        document.add_paragraph(warning, style="List Bullet")

    document.add_paragraph(
        "Nota: la pauta RTC acredita programación publicada; no constituye por sí misma prueba de transmisión efectiva."
    )

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
