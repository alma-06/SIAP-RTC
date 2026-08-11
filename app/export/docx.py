from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.export.contracts import ExportPayload


def export_docx(payload: ExportPayload, output_path: str | Path) -> Path:
    path = Path(output_path)
    document = Document()

    title = document.add_heading(payload.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(payload.subtitle)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("Identificación", level=2)
    identification = document.add_table(rows=2, cols=2)
    identification.cell(0, 0).text = "Periodo"
    identification.cell(0, 1).text = payload.period
    identification.cell(1, 0).text = "Estado"
    identification.cell(1, 1).text = payload.status

    document.add_heading("Indicadores ejecutivos", level=2)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Indicador"
    table.rows[0].cells[1].text = "Valor"
    for label, value in payload.metrics.items():
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)

    document.add_heading("Alertas y consideraciones", level=2)
    if payload.alerts:
        for alert in payload.alerts:
            document.add_paragraph(alert, style="List Bullet")
    else:
        document.add_paragraph("Sin alertas.")

    document.add_heading("Trazabilidad", level=2)
    document.add_paragraph(f"Identificador de evidencia: {payload.evidence_id}")

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)

    document.save(path)
    return path
